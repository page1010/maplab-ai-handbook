from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path("workbook/reviews/JOB-IOS-KOL-INDUSTRY-BRIEF-20260615")
OUT_PATH = OUT_DIR / "AI伺服器供應鏈六大主題_第二層產業分析_20260615.docx"
SCREENSHOT_PATH = Path(
    "/tmp/codex-remote-attachments/019ec6cd-eaf0-7142-99cc-1c9511979ab5/"
    "5a8457d1-24a4-448a-9368-de6e05aba657/1-Photo-1.jpg"
)
TRANSCRIPT_PATH = Path("/private/tmp/ios_kol_u7UPxkcKLSw_full.txt")


SOURCES = [
    (
        "SIA / Deloitte, Powering AI: AI server rack value chain",
        "SIA 2026-06-01 公告指出，AI server rack 內容價值中半導體占比超過 95%，且 AI data center 的資本支出中半導體超過 50%。",
        "https://www.semiconductors.org/new-report-finds-semiconductors-account-for-95-of-an-ai-data-server-racks-value-encompassing-the-full-stack-of-chip-technologies/",
    ),
    (
        "SIA / WSTS April 2026 semiconductor sales",
        "SIA 2026-06-05 公告：2026 年全球半導體銷售預估達 1.5 兆美元，AI infrastructure 與 accelerated computing 是重要驅動。",
        "https://www.semiconductors.org/global-semiconductor-sales-increase-11-month-to-month-in-april/",
    ),
    (
        "TSMC GIGAFAB facilities",
        "TSMC 說明 2024 年 12 吋 GIGAFAB 合計產能超過 1,274 萬片，並在 2024 年建立 2nm 先進製造設施、擴張 Chiayi/Tainan 先進封裝能力以支援 AI 成長。",
        "https://www.tsmc.com/english/dedicatedFoundry/manufacturing/gigafab",
    ),
    (
        "TSMC Advanced Packaging Services",
        "TSMC 說明 3DFabric 包含 SoIC、CoWoS、InFO，並強調 3D silicon stacking 與 advanced packaging 的系統級設計價值。",
        "https://www.tsmc.com/english/dedicatedFoundry/services/advanced-packaging",
    ),
    (
        "TSMC 2026 Q1 quarterly results",
        "TSMC 2026 Q1 guidance 顯示 Q2 revenue guidance 39.0-40.2 billion USD，代表先進製程與 HPC/AI 需求仍需追蹤。",
        "https://investor.tsmc.com/english/quarterly-results/2026/q1",
    ),
    (
        "SEMI equipment forecast reported by Tom's Hardware",
        "SEMI 估計半導體製造設備銷售 2025-2027 連續成長，back-end test 與 assembly/packaging 也受 AI chip 複雜度與 heterogeneous packaging 推動。",
        "https://www.tomshardware.com/tech-industry/semiconductors/sales-of-chip-production-equipment-to-reach-usd156-billion-by-2027-china-taiwan-and-korea-lead-intense-demand",
    ),
    (
        "NVIDIA / ABB 800VDC architecture coverage",
        "公開報導指出 NVIDIA 與 ABB 推進 800VDC power architecture，以支援未來 megawatt-class rack 與 gigawatt-scale AI data center。",
        "https://www.tomshardware.com/tech-industry/big-tech/nvidia-800-vdc-power-rollout-for-1-megawatt-server-racks-to-be-supported-by-abb-company-says-collaboration-will-create-new-power-solutions-for-future-gigawatt-scale-data-centers",
    ),
    (
        "800VDC data center simulation paper",
        "2026 arXiv paper 以 SST-driven 800VDC architecture 模擬 AI data center power delivery，指出 UPS-based AC chain conversion losses 與 fast load transients 是設計痛點。",
        "https://arxiv.org/abs/2601.16502",
    ),
    (
        "NVIDIA CPO / silicon photonics roadmap coverage",
        "公開報導整理 NVIDIA Quantum-X / Spectrum-X Photonics 與 CPO roadmap，重點在更高 bandwidth、更低 per-port power 與更短 electrical trace。",
        "https://www.tomshardware.com/networking/nvidia-outlines-plans-for-using-light-for-communication-between-ai-gpus-by-2026-silicon-photonics-and-co-packaged-optics-may-become-mandatory-for-next-gen-ai-data-centers",
    ),
    (
        "MLCC shortage coverage",
        "公開報導指出 AI server board 與 power supply 對 MLCC 用量提升，部分供應商交期拉長與價格調整，適合作為被動元件緊俏的二級驗證線索。",
        "https://www.pcgamer.com/hardware/guess-what-else-the-pc-industry-is-short-of-now-yes-thats-right-multilayer-ceramic-capacitors/",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")


def set_paragraph_font(paragraph, size: int = 10) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc: Document, text: str = "", size: int = 10, bold: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_bullet(doc: Document, text: str, level: int = 0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=15 if level == 1 else 12, bold=True)
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr_cells[i], "1F4E79")
        set_cell_text(hdr_cells[i], h, bold=True, color="FFFFFF")
        if widths:
            hdr_cells[i].width = Cm(widths[i])

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[i], "F3F6F8")
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    normal.font.size = Pt(10)


def add_source_note(doc: Document) -> None:
    add_heading(doc, "0. 文件來源與使用限制", 1)
    add_bullet(doc, "來源 1：使用者提供之活動截圖，內容包含六大主題：半導體廠務、被動元件、先進封裝、光通訊、功率元件、電子零組件。")
    add_bullet(doc, f"來源 2：YouTube 影片本地 ASR 粗轉錄，檔案位置：{TRANSCRIPT_PATH}。ASR 已確認涵蓋完整影片至約 02:53:33，但專有名詞需人工校正。")
    add_bullet(doc, "來源 3：2026-06-15 即時搜尋之公開資料，優先採用產業協會、公司官方頁、公開技術資料；新聞報導僅作為二級佐證。")
    add_bullet(doc, "限制：本文是產業研究整理與追蹤框架，不是買賣建議；公司名若出現，僅代表供應鏈觀察或影片提及，仍需用財報、法說、出貨認證與客戶導入進度驗證。")


def add_executive_summary(doc: Document) -> None:
    add_heading(doc, "1. 核心結論", 1)
    add_bullet(doc, "第一層需求是 AI GPU/ASIC/HBM 擴張；第二層瓶頸正在往「廠務與無塵室、先進封裝、電源架構、光通訊、被動元件、BBU/散熱/機構」移動。")
    add_bullet(doc, "真正值得追的是供應鏈瓶頸的遷移，而不是只看單一零組件題材。當 GPU 供給改善後，下一個限制會變成能否把高功耗 rack 量產、供電、散熱、連網、備援與維修。")
    add_bullet(doc, "第二層思考的主軸：價值不是平均外溢，而是集中在有設計認證、良率門檻、安規門檻、交期緊張與客戶鎖定能力的供應商。")
    add_bullet(doc, "風險同樣在第二層：CPO/NPO 導入時間延後、800VDC 標準與安規進度、CoWoS/SoIC 擴產節奏、AI capex 放緩、供應商擴產後價格下修。")
    add_bullet(doc, "文件建議把下半年觀察重點放在：CoWoS/SoIC 產能、800V HVDC/正負 400V 過渡架構、25kW BBU、1.6T/3.2T 光通訊、MLCC/牛角電容/電感交期、AI rack 系統 BOM。")


def add_screenshot_rewrite(doc: Document) -> None:
    add_heading(doc, "2. 截圖段落整理", 1)
    rows = [
        ["半導體廠務", "AI 晶片、先進封裝與高階製程擴產，推升半導體廠務工程、無塵室建置與機電整合需求。", "觀察無塵室、機電、氣體化學供應、超純水與電力工程是否成為先進製程放量的前置瓶頸。"],
        ["被動元件", "MLCC、鋁質電容、電感、電阻等需求受 AI 伺服器迭代與 800V HVDC 帶動。", "追蹤高壓、高頻、低 ESR、高溫與長壽命規格，並看 PSU/BBU/rack power 的用量變化。"],
        ["先進封裝", "AI 晶片推動 CoWoS、SoIC、3D 封裝與下一世代異質整合。", "封裝不只是後段製程，而是 AI 晶片能否擴展 memory bandwidth、die size 與 power delivery 的核心。"],
        ["光通訊", "AI data center 高速互聯帶動光模組、NPO/CPO、矽光子供應鏈。", "重點在 optics 逐步靠近 switch ASIC / accelerator，影響既有可插拔模組與新型光引擎價值分配。"],
        ["功率元件", "AI 伺服器電源架構改變、供需失衡與轉單機會形成新一輪功率元件循環。", "觀察 SiC/GaN/MOSFET/整流器/電源模組是否因高壓 DC、SST、BBU 與 rack power 重新定價。"],
        ["電子零組件", "BBU、連接器、散熱、電源、機構件等受 AI server 與 data center 拉動。", "從 server board 擴大到 rack/system BOM，價值轉移到高功率、高可靠、易維修的整機架設計。"],
    ]
    add_table(
        doc,
        ["主題", "截圖段落重點", "第二層整理"],
        rows,
        widths=[3.0, 6.4, 7.1],
    )


def add_transcript_grounding(doc: Document) -> None:
    add_heading(doc, "3. 影片轉檔重點對照", 1)
    add_para(doc, "本節只使用 ASR 作為方向佐證，已將明顯 ASR 錯字做語意校正，例如 CoWoS、SoIC、CPO、HVDC、BBU、MLCC 等。")
    rows = [
        ["開場定位", "主持人指出 AI 影響不只局限在 GPU/伺服器規格，而是延伸到先進製程、先進封裝、電源架構、被動元件、功率元件與光通訊。", "與截圖六大主題完全一致，代表活動主軸是 AI infrastructure supply chain，而非單一 GPU 題材。"],
        ["先進封裝", "逐字稿提到 CoWoS、SoIC、載板/PCB 設備、3D SoIC、Hybrid bonding、CPO 導入等。", "封裝端投資邏輯應從「產能擴張」升級為「先進封裝設備資本密集度與製程精度提高」。"],
        ["800V HVDC / 電源", "逐字稿詳細比較傳統供電、過渡期 HVDC、最終 800V DC 架構，並提到電流降低、線損降低、SST、PowerRack 與 BBU。", "電源架構是 rack 級規格重新設計，不是單一 PSU 瓦數提升；應追蹤正負 400V 到 800V 的遷移路徑。"],
        ["BBU / 電子零組件", "逐字稿提到 BBU 從選配走向標配、5.5kW 到 25kW、SMT 板與被動元件用量增加。", "BBU 是 AI rack 的備援/瞬態電力保險，價值會隨 rack power density 提升。"],
        ["被動元件", "逐字稿承接電源升級，切入 MLCC、鋁質電容、固態電容、電感，以及 PSU/BBU 用量提升。", "被動元件受惠不是只有顆數，而是規格升級、交期、認證與高可靠長壽命設計。"],
        ["光通訊 / CPO", "逐字稿提到玻璃核心載板、Tomahawk 7 CPO 版本、2027-2028 可能量產等觀察。", "CPO/NPO 是下一代資料中心互聯效率議題，但時間點與平台導入節奏需要外部驗證。"],
    ]
    add_table(
        doc,
        ["影片段落", "ASR 抽取重點", "可用判讀"],
        rows,
        widths=[3.2, 7.2, 6.1],
    )


def add_second_order_map(doc: Document) -> None:
    add_heading(doc, "4. 第二層產業地圖", 1)
    rows = [
        ["半導體廠務", "AI/HPC 晶片需求拉動先進製程與封裝廠擴建。", "無塵室、MEP、電力、超純水、氣體化學管線與廠務整合的交期。", "漢唐、聖暉、亞翔、洋基工程等需以接單與毛利驗證。", "新廠建置案、在手訂單、毛利率、工程進度、客戶集中度。", "廠房時程遞延、缺工、原物料、客戶 capex 下修。"],
        ["被動元件", "AI server/rack power 用量上升，電壓/電流/溫度規格升級。", "高可靠 MLCC、鋁電/固態電容、電感與電阻的資格認證與交期。", "國巨、華新科、禾伸堂、鈺邦、立隆電等需分產品線驗證。", "MLCC 交期、ASP、庫存天數、AI/車用/工控占比、PSU/BBU design-in。", "擴產後供過於求、消費電子疲弱、價格競爭。"],
        ["先進封裝", "大 die、HBM、chiplet、memory bandwidth 帶動 CoWoS/SoIC/3D。", "良率、設備精度、載板、測試、臨時鍵合、Hybrid bonding 與週期時間。", "設備、載板、測試、材料廠需以 CoWoS/SoIC 曝險度驗證。", "CoWoS/SoIC 月產能、設備 capex、HBM/ASIC 客戶導入、良率。", "先進封裝節點延後、客戶自研封裝、產能開出後價格壓力。"],
        ["光通訊", "AI cluster scale-out 需要 800G/1.6T/3.2T 互聯。", "從可插拔模組轉向 NPO/CPO/矽光子，價值往光引擎、雷射、封裝與測試移動。", "光模組、雷射、FAU、矽光子封裝與交換器供應鏈。", "NVIDIA/Broadcom switch roadmap、CPO/NPO 設計定案、800G/1.6T 出貨。", "CPO 延後、銅纜/可插拔方案延壽、良率與維修性問題。"],
        ["功率元件", "高功率 rack 需要更高效率電源轉換與保護。", "800VDC / 正負 400V 過渡架構帶動 SiC/GaN/MOSFET、整流、SST、保護元件。", "台達、光寶、群電、功率半導體/模組廠需看規格與客戶認證。", "安規、design win、量產時間、效率、可靠性、毛利率。", "標準不確定、導入時程後移、電源廠垂直整合。"],
        ["電子零組件", "AI server 從單機板變成 rack/system BOM。", "BBU、液冷、連接器、機構件、線材與維修性共同決定可部署容量。", "BBU、散熱、連接器、機構件、ODM/EMS 供應鏈。", "25kW BBU、液冷 CDU/cold plate、connector pin count、rack 出貨。", "客戶規格變動、驗證週期、單一平台過度依賴。"],
    ]
    add_table(
        doc,
        ["主題", "第一層需求", "第二層瓶頸", "觀察名單方向", "驗證指標", "主要風險"],
        rows,
        widths=[2.2, 3.2, 3.8, 3.4, 3.5, 3.2],
    )


def add_deep_dive(doc: Document) -> None:
    add_heading(doc, "5. 主題深挖", 1)
    topics = [
        (
            "5.1 半導體廠務：從蓋廠需求到 capacity unlock",
            [
                "第一層看起來是 AI 晶片需求增加；第二層其實是晶圓廠/封裝廠要能準時開出可用產能。",
                "廠務工程、無塵室、機電整合、超純水、氣體化學管線與電力系統，是先進製程擴張的前置條件。",
                "若無塵室與廠務供給吃緊，受惠者不一定是最大設備商，而是能接大型 turnkey、管理工期與通過客戶安全規範的工程服務商。",
                "關鍵檢查：在手訂單是否真的增加、毛利率是否被缺工/材料侵蝕、工程進度是否與 TSMC/OSAT/記憶體廠 capex 對得上。",
            ],
        ),
        (
            "5.2 先進封裝：不是後段，而是 AI chip scaling 的核心瓶頸",
            [
                "AI processor 需要更高 memory bandwidth、更大 die area 與 chiplet 整合，先進封裝從輔助製程變成架構必要條件。",
                "CoWoS 仍是近期主要 bottleneck，SoIC/Hybrid bonding/3D stacking 則代表更高精度與更高資本密集度。",
                "影片提到的 CoWoS/SoIC/載板/PCB 設備邏輯，應拆成四類追蹤：封裝產能、製程精度、基板/PCB 尺寸升級、測試與良率。",
                "第二層風險是產能擴太快後的價格壓力，以及客戶平台切換造成設備/材料需求節奏改變。",
            ],
        ),
        (
            "5.3 被動元件：用量與規格同時升級",
            [
                "AI server 的被動元件需求不只來自主板，還來自 PSU、BBU、rack power、散熱控制與高速訊號周邊。",
                "MLCC 看顆數與高可靠規格；鋁質/固態/牛角電容看高壓、高溫、低 ESR、壽命與 ripple current；電感看高電流與熱設計。",
                "第二層判斷不是單純『AI server 用更多 MLCC』，而是哪些品項通過客戶認證、交期是否拉長、ASP 是否能維持，以及是否排擠消費電子產能。",
                "若需求真緊，會先反映在交期、價格、稼動率與庫存天數，而不一定立刻反映在營收年增。",
            ],
        ),
        (
            "5.4 光通訊：CPO/NPO 是架構轉移，不是單一產品替換",
            [
                "AI cluster scale-out 讓交換器、NIC、光模組與線材從成本中心變成性能瓶頸。",
                "CPO/NPO 的本質是把光電轉換移近 switch ASIC 或 package，降低 electrical trace loss 與 per-port power。",
                "短期仍要追 800G/1.6T 可插拔模組；中期追 NPO/CPO 導入；長期追矽光子、光引擎、雷射、封裝測試與散熱設計。",
                "最大風險是量產時間表：CPO 的效率價值明確，但維修性、標準化、良率、熱管理與平台採用節奏可能拖慢普及。",
            ],
        ),
        (
            "5.5 功率元件：800VDC 將電源變成 AI infrastructure 主線",
            [
                "高功率 rack 下，傳統多段 AC/DC/AC/DC 轉換的效率、線損、銅用量與散熱壓力會變成設計限制。",
                "800VDC / 正負 400V 是為了降低電流與線損，並支援 megawatt-class rack 和 gigawatt-scale data center。",
                "受惠鏈不只功率半導體，也包含 SST、power shelf、保護元件、connector、BBU、PSU、液冷與整機架驗證。",
                "投資觀察要看實際 design win、安規、量產平台與客戶認證，不要只看『800V』題材字眼。",
            ],
        ),
        (
            "5.6 電子零組件：AI rack system BOM 的價值重分配",
            [
                "AI server 從主機板競爭走向 rack/system competition，BBU、液冷、連接器、機構件與維修設計都會影響可部署容量。",
                "BBU 從選配走向標配的邏輯，是 rack power density 提升後需要應付瞬態負載、短時備援與電網品質問題。",
                "液冷與機構件不是低附加價值配角；當功率密度提升，可靠性、漏液風險、維修時間與客戶認證都會提高進入門檻。",
                "下一步應把電子零組件拆成電力、熱、訊號、機構四條線追蹤，避免把所有零組件混成一個籠統題材。",
            ],
        ),
    ]
    for heading, bullets in topics:
        add_heading(doc, heading, 2)
        for b in bullets:
            add_bullet(doc, b)


def add_tracking_framework(doc: Document) -> None:
    add_heading(doc, "6. 下半年追蹤清單", 1)
    rows = [
        ["產能", "CoWoS/SoIC 月產能、2nm/A16 相關 fab/packaging 進度、OSAT 外溢訂單。", "TSMC 季報/法說、OSAT 法說、設備廠接單。"],
        ["電源", "正負 400V 到 800VDC roadmap、SST、PowerRack、25kW BBU、PSU 瓦數。", "NVIDIA/ODM/電源廠展會、安規、量產時間表。"],
        ["被動", "MLCC/鋁質/固態/電感交期、價格、AI server 用量與庫存天數。", "被動元件廠月營收、法人會議、交期資料。"],
        ["光通訊", "800G/1.6T 出貨、NPO/CPO 採用、矽光子封裝、Tomahawk / Spectrum-X roadmap。", "Broadcom/NVIDIA/Marvell/光模組廠產品更新。"],
        ["廠務", "無塵室/機電/超純水/電力工程 backlog 與毛利。", "工程公司在手訂單、施工進度、客戶 capex。"],
        ["風險", "AI capex 放緩、CPO 延後、800V 標準變動、供應商擴產後價格下跌。", "Hyperscaler capex、庫存、客戶集中度、毛利率轉折。"],
    ]
    add_table(
        doc,
        ["觀察維度", "要追什麼", "可查資料"],
        rows,
        widths=[3.0, 7.3, 6.2],
    )


def add_usable_info(doc: Document) -> None:
    add_heading(doc, "7. 搜尋後可用資訊摘要", 1)
    add_numbered(doc, "SIA/Deloitte 的 AI server rack teardown 可支撐本文主軸：AI 需求會帶動 full-stack semiconductor，而非只帶動 GPU。")
    add_numbered(doc, "TSMC 官方資料可支撐廠務與先進封裝方向：2nm 先進製造、Chiayi/Tainan 先進封裝擴張，以及 3DFabric/CoWoS/SoIC/InFO 系統級封裝。")
    add_numbered(doc, "SEMI 設備市場 forecast 可作為半導體設備與後段封裝設備景氣的二級佐證，但需回頭比對個別台廠營收與產品曝險。")
    add_numbered(doc, "800VDC 相關公開資料支撐電源架構升級，但導入時間可能存在過渡期；文件應用『正負 400V -> 800VDC』路徑追蹤，不宜直接假設全面一次到位。")
    add_numbered(doc, "CPO/NPO 資料顯示光通訊價值會移近 switch ASIC / package，但同時帶來維修、熱、標準與量產時程風險。")
    add_numbered(doc, "MLCC/被動元件報導可當作交期與價格線索，但公司層級仍需由月營收、庫存天數、產品組合與客戶驗證補強。")


def add_source_links(doc: Document) -> None:
    add_heading(doc, "8. 外部資料來源", 1)
    for idx, (title, note, url) in enumerate(SOURCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{idx}. {title}: ")
        set_run_font(run, size=9, bold=True)
        add_hyperlink(p, url, url)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.45)
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run(note)
        set_run_font(run2, size=9)


def add_appendix(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附錄 A. 原始截圖", 1)
    if SCREENSHOT_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(SCREENSHOT_PATH), width=Inches(3.0))
        add_para(doc, f"截圖檔案：{SCREENSHOT_PATH}", size=8)
    else:
        add_para(doc, f"截圖檔案未找到：{SCREENSHOT_PATH}", size=9)
    add_heading(doc, "附錄 B. 本地檔案", 1)
    add_bullet(doc, f"影片 ASR TXT：{TRANSCRIPT_PATH}")
    add_bullet(doc, "/private/tmp/ios_kol_u7UPxkcKLSw_full.srt")
    add_bullet(doc, "/private/tmp/ios_kol_u7UPxkcKLSw_full.json")


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI 伺服器供應鏈六大主題")
    set_run_font(title_run, size=20, bold=True)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("第二層產業分析與追蹤框架")
    set_run_font(subtitle_run, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("整理日期：2026-06-15｜來源：截圖 + YouTube ASR 粗轉錄 + 公開產業資料")
    set_run_font(meta_run, size=9)
    doc.add_paragraph()

    add_source_note(doc)
    add_executive_summary(doc)
    add_screenshot_rewrite(doc)
    add_transcript_grounding(doc)
    add_second_order_map(doc)
    add_deep_dive(doc)
    add_tracking_framework(doc)
    add_usable_info(doc)
    add_source_links(doc)
    add_appendix(doc)

    doc.core_properties.title = "AI 伺服器供應鏈六大主題：第二層產業分析與追蹤框架"
    doc.core_properties.subject = "MAPLAB IOS-KOL industry brief"
    doc.core_properties.author = "MAPLAB IOS-KOL Influencer Radar Manager"
    doc.core_properties.keywords = "AI server, CoWoS, SoIC, 800V HVDC, BBU, CPO, MLCC"

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
