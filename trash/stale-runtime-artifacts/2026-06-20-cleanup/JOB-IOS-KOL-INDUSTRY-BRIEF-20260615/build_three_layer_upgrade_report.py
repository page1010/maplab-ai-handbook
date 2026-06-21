from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("workbook/reviews/JOB-IOS-KOL-INDUSTRY-BRIEF-20260615")
DOCX_OUT = OUT_DIR / "AI伺服器供應鏈六大主題_三層研調整合升級版_20260615.docx"


PUBLIC_SOURCES = [
    (
        "TSMC 4Q25 earnings transcript",
        "https://investor.tsmc.com/english/encrypt/files/encrypt_file/reports/2026-01/51d09df96cd89ac19d65af39032b038dc2896a24/TSMC%204Q25%20Transcript.pdf",
        "2026 capex official range USD52-56B; allocation to advanced process and advanced packaging/testing.",
    ),
    (
        "TSMC 1Q26 earnings transcript",
        "https://investor.tsmc.com/english/encrypt/files/encrypt_file/reports/2026-04/3cef85204275f94fd111485cfdf4adb3c0263c45/TSMC%201Q26%20Transcript.pdf",
        "Capex expected toward high end of USD52-56B; AI demand strong; N3 capacity expansion; advanced packaging capacity very tight.",
    ),
    (
        "TSMC Advanced Packaging Services",
        "https://www.tsmc.com/english/dedicatedFoundry/services/advanced-packaging",
        "3DFabric includes TSMC-SoIC, CoWoS and InFO; collaboration with substrate, memory and materials suppliers.",
    ),
    (
        "SIA/Deloitte AI data-center semiconductor report",
        "https://www.semiconductors.org/powering-ai-the-semiconductor-ecosystem-at-the-foundation-of-data-centers/",
        "AI racks require the full chip stack: logic, memory, analog, foundational chips, networking and power semiconductors.",
    ),
    (
        "arXiv 800 VDC data-center simulation",
        "https://arxiv.org/abs/2601.16502",
        "SST-driven 800 VDC architecture is technically consistent with the power-rack direction, but exact supplier timing remains field intelligence.",
    ),
    (
        "Nexperia official about page",
        "https://www.nexperia.com/about",
        "Official annual revenue shown around USD2.06B and more than 110B units shipped annually; use this to flag the memo's USD75B revenue item as likely unit/currency anomaly.",
    ),
    (
        "PC Gamer / DigiTimes-derived MLCC shortage report",
        "https://www.pcgamer.com/hardware/guess-what-else-the-pc-industry-is-short-of-now-yes-thats-right-multilayer-ceramic-capacitors/",
        "Secondary public source supporting AI-driven MLCC tightness, long lead times and price pressure.",
    ),
]


NUMERIC_CHECKS = [
    [
        "TSMC 2026 capex USD58B; 2027/2028 USD72B/USD84B",
        "Memo L2-L4, L16-L18",
        "2026 official range is USD52-56B and Q1 says toward high end; USD58B is an up-revision expectation, not yet official. 2027/2028 are pure forward field estimates.",
        "第三層前瞻 / 2026部分公開相容",
        "Next TSMC earnings: whether capex range is lifted above USD56B; supplier contract liabilities and fab engineering backlog.",
    ],
    [
        "Arizona fab utility/facility cost 5-6x Taiwan; US fab MEP/cleanroom NTD50-60B per fab",
        "Memo L5-L11",
        "Public sources support overseas fab margin dilution and higher operating cost, but not these exact facility-cost splits.",
        "第三層成本拆解",
        "Track facility-engineering revenue, gross margin, backlog, US project milestones and TSMC overseas margin dilution commentary.",
    ],
    [
        "CoWoS capacity 2027/2028: 190k/255k wafers per month",
        "Memo L28-L29",
        "TSMC publicly confirms advanced packaging is very tight and it works with OSAT partners; exact monthly capacity is not publicly verified.",
        "第三層產能情報",
        "Track TSMC packaging capex mix, OSAT capex, CoWoS equipment orders, HBM layer migration and customer allocation commentary.",
    ],
    [
        "SoIC capacity 2026/2027: 20k/45k wafers per month; CoPoS 2028 10k/month",
        "Memo L30-L36",
        "Publicly consistent with TSMC 3DFabric roadmap and larger die/warpage discussion; exact capacity is not public.",
        "第三層產線情報",
        "Watch hybrid bonding equipment orders, SoIC customer tapeout timing, CoPoS/panel-level pilot-line procurement.",
    ],
    [
        "ASE/SPIL outsourced CoWoS capacity: 25k/month in 2026, 42k in 2027, 50-60k in 2028",
        "Memo L79-L80",
        "TSMC says packaging capacity is tight and OSAT partners are needed; exact outsourced monthly capacity not publicly confirmed.",
        "第三層委外情報",
        "Track OSAT advanced packaging capex, substrate allocation, customer second-source announcements and equipment suppliers' backlog.",
    ],
    [
        "Google TPU Humufish EMIB-T 2027H2; substrate yield 20-50%; 2028 target 10M units",
        "Memo L125-L129",
        "High-value company-call intelligence. Public sources do not validate the exact product, yield or shipment target.",
        "第三層公司call情報",
        "Track ABF substrate capex, EMIB-T wet-process bottlenecks, Google TPU supplier comments and equipment pre-orders.",
    ],
    [
        "GB200/300 racks 26Q1 11k, 26Q2 13k, 26H2 55k-60k; ASIC racks AWS/AMD/Google",
        "Memo L190-L192",
        "Rack units and customer split are not public. They are useful as shipment triangulation anchors.",
        "第三層出貨模型",
        "Compare ODM revenue, power/cooling component orders, BBU shipments, and Nvidia platform transition timing.",
    ],
    [
        "GB200 198kW rack to VR200 444kW; power value USD37k to USD81.2k; Power Rack USD147k; 800V HVDC USD560k",
        "Memo L196-L200, L226-L230",
        "Public sources support direction toward high-density racks and 800VDC; exact rack BOM value is field/BOM intelligence.",
        "第三層BOM模型",
        "Track PSU/BBU ASP, Delta/other power supplier comments, OCP or hyperscaler power architecture releases.",
    ],
    [
        "HVDC efficiency 87.6% to 92.1%; SST 1MW shipment timing, 2.5MW next year, 10MW target",
        "Memo L233-L240",
        "800VDC/SST direction is publicly technically plausible; exact efficiency and supplier timing need company/source validation.",
        "公開相容 + 第三層時程",
        "Track SST certifications, Delta product releases, hyperscaler design wins, and field trial announcements.",
    ],
    [
        "BBU market: 2026 55k racks, 30% penetration, value NTD16.7B; 2027 value NTD30B, YoY +240%",
        "Memo L241-L244",
        "This is an internal model based on rack assumptions and penetration; public data can only verify direction, not the dollar value.",
        "第三層市場模型",
        "Recompute after every ODM/PSU/BBU supplier monthly revenue and platform mix update.",
    ],
    [
        "800G mainstream 2026-2027; 1.6T mass ramp 2027; 3.2T next",
        "Memo L252-L260",
        "Publicly consistent with optical networking upgrade cycle, but supplier-specific OFC observations are field evidence.",
        "公開相容 + 現場觀察",
        "Track OFC/GTC demos, Lumentum/Coherent capex, Nvidia optical purchase agreements and 1.6T module lead times.",
    ],
    [
        "VR200 rack MLCC use around 676k; GB300 around 440k; high-cap 47uF/22uF tight",
        "Memo L312-L327",
        "Public sources support AI-driven MLCC tightness and long lead times; exact rack count is BOM intelligence.",
        "第三層BOM情報",
        "Track ODM BOM revisions, Murata/Samsung/Yageo utilization, lead time, LTA terms and Taiwanese second-source approvals.",
    ],
    [
        "AI server MLCC demand: 2026 40B/month, 2027 100B/month, 2028+ 200B/month",
        "Memo L328-L348",
        "Not publicly verified. Useful as stress-test model because it links demand growth to capacity-expansion lag.",
        "第三層供需模型",
        "Build sensitivity table using rack shipments, MLCC per rack, high-cap capacity multiplier and supplier expansion rate.",
    ],
    [
        "Bullhorn capacitor and sintered foil tightness; PSU/BBU cap counts rise sharply",
        "Memo L374-L410",
        "Mostly company-call / channel intelligence; public verification limited.",
        "第三層供應鏈訊號",
        "Track foil suppliers, Japanese cap maker utilization, PSU platform tear-downs, second-source validation and lead-time change.",
    ],
    [
        "Nexperia revenue scale around USD75B, order spillover USD35B domestic / USD40B international",
        "Memo L421-L423",
        "Public Nexperia official page shows annual revenue around USD2.06B. Treat memo number as likely unit/currency/OCR issue; do not use it until reconciled.",
        "異常 / 需回查",
        "Ask source whether unit is NTD, RMB, cents, or industry TAM rather than Nexperia revenue. Keep transfer-order thesis, reject the exact USD figure for now.",
    ],
    [
        "Power semiconductors: +USD175 value per +1kW; DrMOS 20 -> 32 -> 50+ per chip",
        "Memo L461-L489",
        "Public sources support power-density rise and broader power-chip demand; exact per-kW and DrMOS counts are teardown/BOM intelligence.",
        "第三層BOM模型",
        "Track PSU teardown, PMIC/DrMOS supplier qualification, onsemi/Infineon/Diodes/Taiwan supplier lead times and price actions.",
    ],
]


THEME_SECTIONS = [
    (
        "1. 無塵室與廠務",
        "原判斷：AI晶片與先進封裝擴產使無塵室、機電、氣體、化學、水務與電力系統成為AI供給鏈的底層瓶頸。",
        "升級後：不只看TSMC capex總額，要看工程進場節點、海外成本係數、合約負債與二線工程商外溢。對方的價值在於把fab工程拆到NTD100-120B、US fab NTD500-600B這種可追蹤的廠務拆分。",
        [
            "2026 capex USD58B應標成上修預期；官方目前只到USD52-56B高端。",
            "Arizona/Fab22/Fab20進場時程若提前，廠務供應鏈會比晶圓收入更早反映。",
            "合約負債、在手訂單和工程毛利率比單月營收更早。這是第三層雷達要新增的欄位。",
        ],
    ),
    (
        "2. 設備、先進封裝、載板與PCB",
        "原判斷：CoWoS、SoIC、CoPoS把先進封裝從封測題材升級為半導體設備與材料題材。",
        "升級後：要拆成三個瓶頸：CoWoS既有產能不足、SoIC/hybrid bonding良率與潔淨度、CoPoS/panel-level載具與量測。再往下看ABF/EMIB-T良率與PCB高多層板擴產。",
        [
            "CoWoS 190k/255k、SoIC 20k/45k、CoPoS 10k/month不能當公開事實，但可以當設備訂單的領先假設。",
            "EMIB-T 20-50%良率與Google 10M顆目標是典型公司call edge，價值高，但要和欣興/景碩/臻鼎/IBIDEN/Toppan capex交叉驗證。",
            "玻璃載具/TGV要分清：CoWoS/CoPoS的glass carrier不等於glass core substrate，避免把技術路徑混在一起。",
        ],
    ),
    (
        "3. ODM電子零組件、電源與BBU",
        "原判斷：AI伺服器的供應鏈瓶頸從GPU轉向rack-level power/cooling。",
        "升級後：用平台轉換表來看：GB200/GB300到VR200，不只是晶片功耗增加，而是power shelf、Power Rack、BBU標配化、SST與HVDC架構轉換。",
        [
            "GB200/300與VR200 rack shipment、PSU ASP、BBU滲透率屬第三層BOM模型。",
            "如果BBU由選配走向標配，估值重點不只在電芯，而是BMS、功率模組、電源管理、散熱與安全認證。",
            "800V HVDC會把價值從單顆PSU拉到整個電源櫃；這會改變台達、順達、新盛力、AES等供應鏈的追蹤方法。",
        ],
    ),
    (
        "4. 光通訊",
        "原判斷：AI data center互聯頻寬升級帶動800G、1.6T、3.2T、NPO/CPO/OCS與光引擎需求。",
        "升級後：不要只寫規格升級，要拆成傳輸速率、光源功率、封裝位置、維修性與網路層級。NPO/CPO/OCS不是互斥替代，而是不同距離與網路層級的分工。",
        [
            "800G 2026-2027主流、1.6T 2027放量可保留為公開相容假設。",
            "Lumentum/Coherent/NTT/Google的OFC現場觀察屬前沿情報，要用供應商demo、LTA與量產時程後續驗證。",
            "CPO帶來的不是傳統光模組消失，而是光源、socket、外部雷射、測試與維修模型改變。",
        ],
    ),
    (
        "5. 被動元件",
        "原判斷：AI伺服器功率提升帶動MLCC、鋁電容、電感、電阻需求。",
        "升級後：對方的強項是BOM數字化，讓被動元件從概念題變成供需模型。VR200 rack MLCC約676k、47uF/22uF高容值吃產能3-7倍，是判斷緊缺的核心。",
        [
            "MLCC公開資料能驗到AI造成緊缺、交期延長與漲價，但不能驗到每櫃676k與40B/100B/month需求。",
            "第三層追蹤要新增：高容值料號、產能消耗倍率、LTA、second source、ODM推遲漲價能力。",
            "鋁電容/牛角電容/燒結箔的缺口比一般MLCC更接近供應鏈早期訊號，應列入下次call問題。",
        ],
    ),
    (
        "6. 功率元件",
        "原判斷：PSU瓦數、800V HVDC、SiC/GaN/MOSFET/DrMOS用量提升，會開啟功率元件上升循環。",
        "升級後：保留功率上升主線，但必須修正Nexperia數字。Nexperia轉單與Yangjie制裁是有效催化，USD75B revenue數字與公開Nexperia annual revenue不一致，需回查單位。",
        [
            "不要因一個數字異常就丟掉整段研調；應拆成「事件方向有效」與「金額單位待查」。",
            "Infineon/Onsemi/Diodes/Taiwan suppliers轉單要用料號重疊、Hot Run、第二供應商資格與漲價公告驗證。",
            "+USD175/kW、DrMOS 20->32->50+是BOM模型，之後應用teardown和供應商法說交叉驗證。",
        ],
    ),
]


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8.5)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[idx], "F7F9FB")
            set_cell_text(cells[idx], value)
    set_table_width(table, widths)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4E79", 14, 7),
        ("Heading 2", 13, "1F4E79", 10, 5),
        ("Heading 3", 11.5, "1F4E79", 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.08


def add_para(doc: Document, text: str, *, bold: bool = False, size: float = 10.5, color: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=10)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run, size=10)


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color="1F4E79")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.5)
    set_table_width(table, [6.8])
    doc.add_paragraph()


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 伺服器供應鏈六大主題：三層研調整合升級版")
    set_run_font(r, size=18, bold=True, color="1F4E79")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("截圖 + 影片逐字稿 + 定錨產業筆記 MEMO(5) + 公開驗證")
    set_run_font(r2, size=11.5, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("整理日期：2026-06-15｜角色：MAPLAB IOS-KOL Influencer Radar Manager")
    set_run_font(r3, size=9.5, color="555555")


def build_report() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "版本結論",
        "我的主線判斷不推翻，但權重上修：廠務/先進封裝設備/被動元件/功率元件比原稿更重要。"
        "對方的edge在於公司call與BOM拆解，尤其第三層數字應保留為領先訊號；公開驗證的責任是標註邊界，而不是把不能公開驗證的情報刪掉。",
    )

    doc.add_heading("1. 三層判讀框架", level=1)
    add_para(
        doc,
        "本版把資訊分成三層，避免把公開資料、研究員公司call、我們自己的推論混成同一種確定性。",
    )
    add_table(
        doc,
        ["層級", "定義", "適合放什麼", "輸出規則"],
        [
            [
                "第一層：公開敘事",
                "可由公司法說、官網、公開報告、新聞或產業會議驗到的方向。",
                "AI/HPC需求、TSMC capex區間、先進封裝緊、800VDC方向、MLCC供應吃緊。",
                "可以寫成事實，但要附來源與日期。",
            ],
            [
                "第二層：產業結構",
                "從公開事實往供需、瓶頸、良率、認證、資本支出與議價權推導。",
                "GPU瓶頸轉為rack power/cooling、SoIC推升CoWoS而非取代、BBU標配化。",
                "要寫成因果鏈，並列出反證條件。",
            ],
            [
                "第三層：春江水暖情報",
                "產業研究員跑公司、call公司、看展會、拆BOM、問供應鏈得到的早期訊號。",
                "月產能、單櫃用量、良率、客戶出貨、capex拆分、轉單與second source。",
                "保留數字，但標記為edge intelligence；用後續公開數據追蹤，不直接當官方事實。",
            ],
        ],
        [1.1, 2.0, 2.1, 1.6],
    )

    doc.add_heading("2. 對方數字的整合與查核", level=1)
    add_para(
        doc,
        "下表不是為了否定對方，而是把他的數字變成可追蹤的研究資產。分類規則："
        "「公開可驗」代表可直接引用；「公開相容」代表方向能被公開資料支持但精確數字未驗；"
        "「第三層」代表公司call/BOM/供應鏈情報；「異常」代表公開資料衝突，需回問單位或口徑。",
    )
    add_table(
        doc,
        ["數字 / 觀點", "來源位置", "查核結論", "分類", "後續驗證點"],
        NUMERIC_CHECKS,
        [1.55, 0.75, 2.25, 1.0, 1.65],
    )

    doc.add_heading("3. 六大主題：升級後的第三層思考", level=1)
    for heading, original, upgrade, bullets in THEME_SECTIONS:
        doc.add_heading(heading, level=2)
        add_para(doc, original, bold=True)
        add_para(doc, upgrade)
        for item in bullets:
            add_bullet(doc, item)

    doc.add_heading("4. 我方格式升級：把整理能力往研調員方向強化", level=1)
    add_para(
        doc,
        "原本我方優勢是結構清楚、能把截圖與逐字稿快速整理成主題地圖；短板是數字密度不足，第三層情報沒有被充分保留下來。"
        "升級後，IOS-KOL的標準格式要同時包含結構化、數字表、確定性標籤、驗證路徑與反證條件。",
    )
    for item in [
        "每個主題至少要有一個「第三層雷達訊號」：如良率、月產能、單櫃用量、second source、LTA、capex、設備交期、Hot Run。",
        "每個數字都要標籤：官方公開、公開相容、公司call情報、BOM模型、異常待查。",
        "對不能公開驗證的數字，不刪掉；改寫成「領先假設」並列公開追蹤條件。",
        "遇到數字與公開資料衝突時，拆成方向與數字兩部分：方向可能成立，數字口徑需回查。",
        "文件末尾固定留下下一次研究問題，讓團隊能拿去call公司或追月營收，而不是只停在文字摘要。",
    ]:
        add_numbered(doc, item)

    doc.add_heading("5. 下次call公司 / 追資料問題清單", level=1)
    add_table(
        doc,
        ["主題", "應問的第三層問題", "用來驗證什麼"],
        [
            ["廠務", "2026Q3後哪些fab/site進場？合約負債是否墊高？美國案毛利與付款節點？", "TSMC capex上修與fab工程外溢。"],
            ["設備/封裝", "SoIC/hybrid bonding設備占比何時顯著？CoPoS是否2027開始建線？", "SoIC/CoPoS提前的可能性。"],
            ["載板/PCB", "EMIB-T良率、濕製程設備交期、ABF客戶是否預訂2028產能？", "Google TPU/ASIC拉貨是否會吃掉載板產能。"],
            ["電源/BBU", "BBU由選配到標配的客戶節點？Power Rack ASP與滲透率？", "BBU市場NTD16.7B/30B模型。"],
            ["光通訊", "1.6T客戶拉貨時間、CPO/NPO/OCS採用層級、ELS雷射功率規格？", "800G->1.6T->3.2T升級曲線。"],
            ["被動", "47uF/22uF交期、LTA價格、second source核准、燒結箔供給？", "MLCC與牛角電容是否真緊缺。"],
            ["功率", "Nexperia金額口徑、Infineon退出哪些成熟品、台系料號是否完成Hot Run？", "轉單效應與台系替代料機會。"],
        ],
        [1.0, 3.55, 2.25],
    )

    doc.add_heading("6. 公開來源與本版邊界", level=1)
    add_para(
        doc,
        "本版使用公開來源驗證方向與邊界；對方PDF中的公司call、BOM、良率與產能數字多數無法由公開資料直接驗證，已用第三層標籤保留。"
    )
    add_table(
        doc,
        ["來源", "URL", "本版使用方式"],
        PUBLIC_SOURCES,
        [1.7, 3.05, 2.05],
    )
    add_para(
        doc,
        "投資判斷提醒：本文件用於產業研究與KOL雷達，不構成投資建議。所有第三層情報都需要以公司法說、月營收、出貨、客戶認證與價格/交期變化持續驗證。",
        bold=True,
    )

    doc.core_properties.title = "AI伺服器供應鏈六大主題：三層研調整合升級版"
    doc.core_properties.author = "MAPLAB IOS-KOL Influencer Radar Manager"
    doc.core_properties.subject = "IOS-KOL third-layer industry research upgrade"
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build_report()
