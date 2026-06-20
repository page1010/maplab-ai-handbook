from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("workbook/reviews/JOB-IOS-KOL-INDUSTRY-BRIEF-20260615")
DOCX_OUT = OUT_DIR / "定錨2026年中產業筆記_AI供應鏈研究報告_重點推演推薦_20260615.docx"
MD_OUT = OUT_DIR / "定錨2026年中產業筆記_AI供應鏈研究報告_重點推演推薦_20260615.md"


THEMES = [
    {
        "name": "無塵室與廠務工程",
        "priority": "A",
        "pdf": [
            "TSMC 2026 capex 預估 USD58B，2027/2028 進一步到 USD72B/USD84B；A14/A10 帶動高原期。",
            "Arizona 廠務成本較台灣大幅提高，且嫡系廠商外溢到二線工程商。",
            "Fab21、Fab22、Fab20 多個主系統裝機節點提前，Q3 起工程合約負債可能墊高。",
        ],
        "second": "AI 對先進製程與先進封裝需求不是只拉晶圓代工收入，會先拉動建廠、無塵室、機電、水氣化工程。廠務公司收入確認通常早於晶圓量產，且海外 fab 的單座工程價值量更高。",
        "third": "PDF 的 edge 在於把 capex、Fab 進場節點與單座 fab 工程拆分成可追蹤數字。這些數字尚非官方公開確認，但能用合約負債、在手訂單、工程進場、二線廠商接單來驗證。",
        "why": "推薦追蹤的原因是能見度高、前置性強，且二線外溢代表供給吃緊程度高於市場只看龍頭工程商的理解。優先看已在 TSMC 供應鏈且出現海外/先進製程訂單外溢的廠務與機電工程商。",
        "watch": "TSMC 是否正式上修 capex；Fab22/Fab21 進場是否如期；工程商合約負債與毛利是否同步上升。",
    },
    {
        "name": "半導體設備與先進封裝",
        "priority": "A",
        "pdf": [
            "CoWoS 2027/2028 年底月產能估 190k/255k。",
            "SoIC 2026/2027 月產能估 20k/45k；CoPoS 2028 估 10k/month。",
            "SoIC/CoPoS 設備資本密集度約較 CoWoS 高 50%，製程含 CMP、濕製程、pick and place、bonding、熱製程。",
        ],
        "second": "AI 晶片從單 die 走向 chiplet 與更大 package，瓶頸由晶圓製造延伸到 hybrid bonding、RDL、量測、清洗、載具與封裝良率。設備商受惠不只是 CoWoS 擴產，還包含 SoIC 與 CoPoS 的站點增量。",
        "third": "PDF 給出各封裝技術的月產能節奏與良率/潔淨度要求，這是公司 call 型數字。它讓我們可以從『概念受惠』升級為『哪些製程站點、哪些設備商、哪一年開始營收占比提升』。",
        "why": "推薦追蹤已切入關鍵站點且可吃到多技術路徑的設備商，例如濕製程、CMP、bonding、檢測、RDL/PVD、panel-level 封裝設備。這類廠商比單一終端需求更像 picks-and-shovels。",
        "watch": "CoWoS 設備訂單是否延續；SoIC/hybrid bonding 站點採購是否在 2027 明顯放大；CoPoS 是否提前建置第一條線。",
    },
    {
        "name": "ABF載板、EMIB-T與PCB設備",
        "priority": "A",
        "pdf": [
            "Google TPU Humufish 2027H2 改採 EMIB-T，載板段良率約 20-50%，2028 目標 10M 顆。",
            "欣興 capex 接近 NTD500B，光復/楊梅多廠區擴產；景碩、臻鼎、IBIDEN、Toppan 同步擴。",
            "PCB 板廠 capex 倍增，AI 主板 HLC/HDI/mSAP 規格提升；牧德、創鋒、揚博、群翊、尖點、大量等設備受惠。",
        ],
        "second": "ASIC 與 GPU 封裝尺寸變大、層數提高、良率壓力增加，會消耗更多 ABF 與 PCB 產能。若 EMIB-T 良率低，需求不是線性增加，而是同樣出貨量需要更多載板與設備備援。",
        "third": "Humufish/EMIB-T 20-50% 良率、Google 10M 顆目標與廠區進設備時程，是典型第三層 field intelligence。這些數字能領先公開財報，指向設備 pre-order 與 capex 提前。",
        "why": "推薦設備鏈優先於純載板廠，因為設備訂單會先反映擴產，且可同時吃 EMIB-T、ABF、CoWoS、PCB 高多層板多條需求線。載板廠則需追蹤 capex 轉稼動率與折舊壓力。",
        "watch": "欣興/臻鼎/景碩 capex 是否續上修；PCB 設備商訂單能見度是否延到 2027；EMIB-T 良率是否改善或反而擴大載板消耗。",
    },
    {
        "name": "Power Rack、BBU與散熱零組件",
        "priority": "A",
        "pdf": [
            "GB200/GB300 到 VR200，rack 功耗由 198kW 到 444kW，VR200 power shelf 由 6 顆 18.3kW PSU 構成。",
            "GB200 power 價值 USD37k，VR200 提升至 USD81.2k；Power Rack USD147k；800V HVDC power rack 可到 USD560k。",
            "BBU 由選配轉標配；2026 55k 櫃、30% 滲透率、產值約 NTD16.7B；2027 約 NTD30B，YoY +240%。",
        ],
        "second": "AI rack 功率密度提升後，價值不只在伺服器內 PSU，而是上移到 Power Rack、BBU、HVDC、SST、備援控制與能源管理。這會讓電源供應鏈從零組件升級為系統級解決方案。",
        "third": "rack shipment、BBU 滲透率、Power Rack ASP、SST 時程都屬 BOM/field model。它們無法從公開資料直接驗證，但可用台達供應商、BBU 出貨、PSU 功率規格與安規認證交叉驗證。",
        "why": "推薦追蹤 power/BBU 供應鏈，因為單櫃價值量提升倍數大，且 HVDC 會重寫供應鏈分工。台達、順達、新盛力、AES 等供應商應以平台別 ASP 與滲透率追蹤。",
        "watch": "800V HVDC 導入節點；BBU 是否由 NVIDIA/AWS TR3 擴到更多 CSP；SST 1MW/2.5MW/10MW 產品驗證與量產時程。",
    },
    {
        "name": "光通訊與光引擎",
        "priority": "B+",
        "pdf": [
            "800G 為 2026-2027 主流；1.6T 大規模放量落在 2027；3.2T 為下一世代。",
            "EML、CW Laser、DSP、PD、先進封裝測試價值量同步提升。",
            "OFC 現場觀察：Lumentum 200G/400G EML、高功率 CW Laser；Coherent/NTT NPO socket；OCS 技術路徑分化。",
        ],
        "second": "GPU scale-up/scale-out 需要更高頻寬與更低功耗互聯，光通訊規格升級速度快於供應鏈消化能力。真正受惠點不是只看模組量，而是高階光源、DSP、socket、測試與 CPO/NPO/OCS 轉換。",
        "third": "OFC 展場 demo、特定 socket 供應商與 Google/Lumentum/Coherent 技術路徑屬前沿觀察。這些訊號能提前判斷誰會在 1.6T/3.2T、NPO/CPO/OCS 分工中取得位置。",
        "why": "推薦列入中期追蹤，因為規格升級確定，但量產節點與贏家分化仍需驗證。優先看同時具備高階 EML/CW Laser、InP 能力、socket/OCS 入口或資料中心交換器客戶的供應商。",
        "watch": "1.6T 訂單是否在 2027 放量；CPO/NPO 是否進入量產設計；OCS 從 Spine/DCI 滲透到更靠近 GPU 的時間。",
    },
    {
        "name": "被動元件：MLCC、鋁電容、電感",
        "priority": "A",
        "pdf": [
            "VR200 整櫃 MLCC 約 676k 顆，GB300 約 440k 顆；功耗與 MLCC 用量約同步成長。",
            "AI server MLCC 需求估 2026 40B/month、2027 100B/month、後年 200B/month；高容值 MLCC 產能消耗 3-7 倍。",
            "牛角電容、Hybrid、SP-CAP、Polymer 與功率電感同步受惠；燒結箔與高容值料號供需緊缺。",
        ],
        "second": "P=VI 下功率提升會帶動電容/電感的穩壓與儲能需求。AI server 高功率、高電流與高密度架構使被動元件從低估值成熟品變成受限料號，ASP、交期、second source 重要性提高。",
        "third": "每櫃 MLCC 顆數、47uF/22uF 料號、PSU/BBU 牛角電容數量、燒結箔供給，是最有價值的第三層 BOM 情報。它能讓我們從『被動元件受惠』推到『哪個料號最缺、誰能二供』。",
        "why": "推薦追蹤高容值 MLCC、AI 用鋁電容與功率電感供應鏈。優先看已切入 VR200/電源端、具備 second source 角色、且可受惠交期延長或漲價的台廠。",
        "watch": "Murata/Samsung/Yageo 擴產速度；ODM LTA 是否壓抑漲價；日廠 OS-CON/SP-CAP 減產或轉單；燒結箔供給是否緩解。",
    },
    {
        "name": "功率元件、DrMOS與供應鏈重組",
        "priority": "A-",
        "pdf": [
            "Nexperia 退出/供給受限與揚杰制裁推動替代料 hot run，轉單效應開始發酵。",
            "AI PSU 從 5kW/8kW 走向 12kW/18kW，多相 PFC/LLC、SiC、GaN、MOSFET 用量增加。",
            "單櫃每提升 1kW 約帶動功率元件 USD175；DrMOS 從 700W 約 20 顆、1300W 約 32 顆，到 1800W 以上 50+ 顆。",
        ],
        "second": "高功率密度與 800V HVDC 使功率元件用量與規格同步升級；同時國際 IDM 把產能轉向 SiC/SST/高毛利 AI 產品，成熟 MOSFET/二極體/TVS 供給外溢給台系二供。",
        "third": "PDF 的 Nexperia USD75B 數字需標 anomaly，但替代料 hot run、DrMOS 顆數與每 kW 價值量仍是有用的供應鏈訊號。應保留轉單 thesis，剔除未核對的營收量級。",
        "why": "推薦追蹤具產品重疊與自有製造/封裝能力的台系功率廠，以及已推出 DrMOS、MOS PMIC、AI server 替代料者。這是『AI需求拉高』與『地緣供應鏈重組』雙催化。",
        "watch": "Nexperia/揚杰替代料是否正式量產；Infineon/onsemi 是否繼續轉高階；台系 DrMOS/PMIC 驗證節點；價格調漲是否擴散。",
    },
]


NUMERIC_MATRIX = [
    ["無塵室", "TSMC 2026 capex USD58B；2027/2028 USD72B/USD84B", "company_call / public_compatible", "2026 官方公開仍以 USD52-56B 高端為基準；58B 是上修訊號。", "若官方不上修，廠務彈性會下降。"],
    ["廠務", "美國單座 fab 無塵室+機電 NTD500-600B；台灣 NTD100-120B", "company_call", "公開資料支持海外成本較高，但精確拆分未公開。", "需看工程商 backlog 與毛利。"],
    ["CoWoS", "2027/2028 月產能 190k/255k", "company_call", "公開方向相容；TSMC 確認先進封裝很緊，但未公開精確月產能。", "客戶導入/設備交期若延後，產能節奏下修。"],
    ["SoIC/CoPoS", "SoIC 20k/45k；CoPoS 2028 10k/month；設備資本密集度 +50%", "company_call", "技術方向相容；精確數字待驗。", "Feynman non-SoIC 或良率未達標會延後。"],
    ["EMIB-T/ABF", "Humufish EMIB-T 良率 20-50%；2028 10M 顆", "company_call", "高價值第三層情報，公開未驗。", "良率若快速改善，載板消耗強度會下降。"],
    ["Power/BBU", "VR200 power value USD81.2k；Power Rack USD147k；HVDC USD560k", "bom_model", "方向與高功率 rack 一致；金額需供應鏈驗證。", "HVDC 導入慢於預期會推遲價值量。"],
    ["BBU", "2026 55k rack、30%滲透、NTD16.7B；2027 NTD30B", "bom_model", "模型依 rack 出貨與滲透率，非公開事實。", "rack 出貨或 BBU 標配化不如預期。"],
    ["光通訊", "800G 2026-2027 主流；1.6T 2027 放量；3.2T 下一代", "public_compatible", "與產業規格升級相符，但供應商贏家仍需驗證。", "CPO/NPO 採用慢、價格壓力。"],
    ["MLCC", "VR200 rack 約 676k 顆；AI server 需求 40B/100B/200B per month", "bom_model / channel_check", "公開支撐 MLCC 緊缺，精確需求模型待驗。", "ODM LTA 壓價、擴產快於需求。"],
    ["鋁電容", "18kW PSU 牛角 16-18 顆；25kW 20+ 顆；燒結箔 tight", "bom_model / channel_check", "主要是 tear-down 與供應鏈口徑。", "second source 導入失敗或日廠擴產。"],
    ["功率元件", "每 +1kW 帶動 USD175；DrMOS 20 -> 32 -> 50+", "bom_model", "方向相容，精確數字需 teardown 驗證。", "台系驗證不過或 IDM 供給回補。"],
    ["Nexperia", "PDF 稱 USD75B revenue/order scale", "anomaly", "公開 Nexperia revenue 量級約 USD2B，PDF 數字需核對單位/期間/TAM。", "保留轉單 thesis，但不採用 USD75B 作事實。"],
]


PRIORITY_ROWS = [
    ["A", "先進封裝設備", "CoWoS 續擴 + SoIC/CoPoS 資本密集度更高", "設備訂單領先收入、站點多、可吃多技術路徑", "弘塑、辛耘、志聖、均華、萬潤、倍利科、天虹、由田、政美應用等"],
    ["A", "廠務/無塵室", "TSMC capex 高原期 + 海外 fab 單座價值量提高", "合約負債與二線外溢是早期訊號", "漢唐、和淞、兆聯、帆宣、亞翔、洋基工程等"],
    ["A", "ABF/PCB設備", "EMIB-T 良率低 + AI板層數/面積增加", "設備供應鏈先吃擴產，彈性可能高於載板製造端", "敘豐、牧德、創鋒、揚博、群翊、尖點、大量、盟立、長廣等"],
    ["A", "被動元件", "VR200 BOM 數字化揭示高容值 MLCC/鋁電容供需缺口", "料號緊缺、second source、漲價/交期都是可驗證 edge", "國巨、華新科、禾伸堂、立隆電、鈺邦、金山電、臺慶科、信昌電等"],
    ["A", "Power Rack/BBU", "Power shelf -> Power Rack -> HVDC，單櫃價值量跳升", "BBU 標配化與 SST/HVDC 時程是第三層雷達", "台達、順達、新盛力、AES 及相關 BMS/功率供應鏈"],
    ["A-", "功率元件", "AI功率密度提升 + Nexperia/揚杰供應鏈重組", "有價格、轉單、替代料 hot run 三個催化", "台半、強茂、德微、富鼎、力智、立錡、茂達、矽力-KY等"],
    ["B+", "光通訊", "800G/1.6T/3.2T 升級確定，NPO/CPO/OCS 分工未定", "中期大方向強，但贏家與量產節點需驗證", "Lumentum、Coherent、聯亞、環宇-KY、全新、穩懋、IET-KY、嘉基、鴻海、智邦、啟碁等"],
]


FOLLOW_UP_QUESTIONS = [
    "TSMC 下一次法說是否把 2026 capex 區間正式上移到 USD54-58B 或高於 USD56B？",
    "Fab22 P4/P5、Fab21 P3/P4、Fab20 P3 的主系統進場是否照 PDF 提前時程發生？",
    "SoIC hybrid bonding 設備訂單是否在 2027 明顯放大？CoPoS 第一條線是否提前到 2027？",
    "Google Humufish EMIB-T 良率是否仍在 20-50%？ABF/PCB 設備 pre-order 是否開始反映 2028 10M 顆目標？",
    "VR200 rack power/cooling BOM 是否確認：Power Rack、BBU、HVDC、SST 的 ASP 與滲透率如何？",
    "47uF/22uF MLCC、燒結箔、SP-CAP、Hybrid 電容的 lead time 與 price 是否延續？台廠 second source 驗證到哪一步？",
    "Nexperia USD75B 數字到底是幣別、期間、TAM 還是 OCR/單位錯誤？轉單 thesis 是否有客戶 hot-run 佐證？",
]


PUBLIC_CHECKS = [
    "TSMC 官方公開資料可支撐 2026 capex USD52-56B 且往高端，但 PDF 的 USD58B 是上修期待，需標記為第三層。",
    "TSMC 3DFabric/先進封裝方向與 CoWoS/SoIC/advanced packaging tightness 相容，但月產能精確數字仍屬公司 call。",
    "Nexperia 官方營收量級與 PDF 的 USD75B 明顯不一致，應列 anomaly，不可直接採用。",
    "800V HVDC、SST、AI data center power-density 方向具公開技術合理性，但 rack BOM 價值量仍需以供應鏈與 teardown 驗證。",
    "MLCC 緊缺、AI server 拉貨與交期延長已有公開報導支持，但每櫃顆數與月需求模型仍屬第三層 BOM 情報。",
]


SOURCE_LINKS = [
    "PDF：2026 年中產業趨勢講座 / 定錨產業筆記 / 2026.06.12 / 原始 PDF：dingmao_2026_midyear_industry_memo.pdf",
    "TSMC 1Q26 transcript：官方 capex 區間、AI/HPC 與先進封裝 tightness。",
    "TSMC Advanced Packaging Services：CoWoS、SoIC、InFO/3DFabric 方向。",
    "Nexperia official about：營收量級用於校正 PDF anomaly。",
    "SIA/Deloitte AI data-center semiconductor ecosystem：AI data center 對完整半導體/電源/網通供應鏈需求的公開背景。",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc: Document, headers, rows, widths, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    repeat_table_header(table.rows[0])
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(31, 77, 120)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.size = Pt(font_size)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items, style="List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill="F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    set_table_widths(table, [9360])
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1


def build_markdown() -> str:
    lines = [
        "# 定錨 2026 年中產業筆記：AI 供應鏈研究報告",
        "",
        "資料基礎：原始 PDF《2026 年中產業趨勢講座 / 定錨產業筆記 / 2026.06.12》。",
        "",
        "## 讀法",
        "",
        "本報告的「推薦」定義為研究優先級與追蹤名單，不是個人化買賣建議。所有第三層數字都保留其 edge，但不寫成官方事實。",
        "",
        "## 核心結論",
        "",
        "- A 級主線：先進封裝設備、廠務、ABF/PCB設備、被動元件、Power Rack/BBU。",
        "- A- 主線：功率元件，受惠 AI 功率密度與供應鏈重組，但 Nexperia USD75B 數字需回查。",
        "- B+ 主線：光通訊，方向明確，但 1.6T/3.2T、NPO/CPO/OCS 贏家與量產節點仍需驗證。",
        "",
        "## 推薦優先級",
        "",
        "| 優先級 | 主線 | 為什麼受惠 | 為什麼推薦/追蹤 | PDF點名族群 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in PRIORITY_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.extend(["", "## 六大主題重點與二三層推演", ""])
    for theme in THEMES:
        lines.extend(
            [
                f"### {theme['priority']}｜{theme['name']}",
                "",
                "**PDF重點**",
                "",
            ]
        )
        lines.extend([f"- {item}" for item in theme["pdf"]])
        lines.extend(
            [
                "",
                f"**第二層推演**：{theme['second']}",
                "",
                f"**第三層edge**：{theme['third']}",
                "",
                f"**為什麼推薦/追蹤**：{theme['why']}",
                "",
                f"**反證/驗證**：{theme['watch']}",
                "",
            ]
        )
    lines.extend(
        [
            "## 數字驗證矩陣",
            "",
            "| 主題 | PDF數字 | Tag | 採用方式 | 風險/反證 |",
            "| --- | --- | --- | --- | --- |",
        ]
    )
    for row in NUMERIC_MATRIX:
        lines.append("| " + " | ".join(row) + " |")
    lines.extend(["", "## 公開驗證備註", ""])
    lines.extend([f"- {item}" for item in PUBLIC_CHECKS])
    lines.extend(["", "## 下次 call 問題", ""])
    lines.extend([f"- {item}" for item in FOLLOW_UP_QUESTIONS])
    lines.extend(["", "## 來源", ""])
    lines.extend([f"- {item}" for item in SOURCE_LINKS])
    return "\n".join(lines) + "\n"


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("定錨 2026 年中產業筆記：AI 供應鏈研究報告")
    r.font.size = Pt(21)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sub.add_run("原始 PDF 重點標註、第二/第三層受惠推演、推薦追蹤邏輯 | 2026-06-15").italic = True

    add_callout(
        doc,
        "閱讀定位",
        "本報告的「推薦」是研究優先級與追蹤名單，不是個人化買賣建議。PDF 中的公司 call、BOM、rack、良率、產能數字保留其資訊 edge，但一律標記為第三層或模型數字，需後續驗證。",
        fill="E8EEF5",
    )

    doc.add_heading("一頁結論", level=1)
    add_bullets(
        doc,
        [
            "A 級主線：先進封裝設備、廠務、ABF/PCB設備、被動元件、Power Rack/BBU。",
            "A- 主線：功率元件，受惠 AI 功率密度與供應鏈重組，但 Nexperia USD75B 數字需回查。",
            "B+ 主線：光通訊，方向明確，但 1.6T/3.2T、NPO/CPO/OCS 贏家與量產節點仍需驗證。",
            "核心方法：PDF重點不是直接當事實，而是變成可追蹤的 call 問題、BOM敏感度與供應鏈雷達。",
        ],
    )

    doc.add_heading("推薦優先級地圖", level=1)
    add_table(
        doc,
        ["級別", "主線", "為什麼受惠", "為什麼推薦/追蹤", "PDF點名族群"],
        PRIORITY_ROWS,
        [650, 1280, 2450, 2450, 2530],
        font_size=8.1,
    )

    doc.add_heading("受惠傳導圖", level=1)
    add_callout(
        doc,
        "AI算力需求 -> 先進製程/封裝 -> 廠務/設備/載板 -> rack power/cooling -> 光通訊/被動/功率元件",
        "第二層的核心是供應鏈瓶頸遞延：GPU需求先拉動TSMC capex與封裝產能，再拉動廠務、設備、ABF/PCB、power rack、散熱、光通訊、被動元件與功率半導體。第三層的價值是把每一段用產能、良率、BOM、ASP、滲透率與交期數字化。",
        fill="F4F6F9",
    )

    doc.add_heading("六大主題重點與二三層推演", level=1)
    for theme in THEMES:
        doc.add_heading(f"{theme['priority']}｜{theme['name']}", level=2)
        doc.add_paragraph("PDF重點", style="Heading 3")
        add_bullets(doc, theme["pdf"])
        add_callout(doc, "第二層：為什麼受惠", theme["second"], fill="F7FAFC")
        add_callout(doc, "第三層：PDF的edge在哪裡", theme["third"], fill="FFF8E8")
        add_callout(doc, "為什麼推薦/追蹤", theme["why"], fill="EEF7EE")
        p = doc.add_paragraph()
        p.add_run("反證/驗證：").bold = True
        p.add_run(theme["watch"])

    doc.add_heading("數字驗證矩陣", level=1)
    add_table(
        doc,
        ["主題", "PDF數字", "Tag", "採用方式", "風險/反證"],
        NUMERIC_MATRIX,
        [900, 2800, 1450, 2550, 1660],
        font_size=7.7,
    )

    doc.add_heading("公開驗證備註", level=1)
    add_bullets(doc, PUBLIC_CHECKS)

    doc.add_heading("下次 call 問題", level=1)
    add_bullets(doc, FOLLOW_UP_QUESTIONS, style="List Number")

    doc.add_heading("來源", level=1)
    add_bullets(doc, SOURCE_LINKS)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("IOS-KOL research brief | third-layer tags required")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(DOCX_OUT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print({"docx": str(DOCX_OUT), "md": str(MD_OUT)})


if __name__ == "__main__":
    main()
