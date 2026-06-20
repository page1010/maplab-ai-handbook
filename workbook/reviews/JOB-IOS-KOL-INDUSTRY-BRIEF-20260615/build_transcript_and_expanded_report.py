from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path("workbook/reviews/JOB-IOS-KOL-INDUSTRY-BRIEF-20260615")
SRT_PATH = Path("/private/tmp/ios_kol_u7UPxkcKLSw_full.srt")
TXT_PATH = Path("/private/tmp/ios_kol_u7UPxkcKLSw_full.txt")
SCREENSHOT_PATH = Path(
    "/tmp/codex-remote-attachments/019ec6cd-eaf0-7142-99cc-1c9511979ab5/"
    "5a8457d1-24a4-448a-9368-de6e05aba657/1-Photo-1.jpg"
)

TRANSCRIPT_TXT_OUT = OUT_DIR / "IOS-KOL影片完整逐字稿_含時間碼_20260615.txt"
TRANSCRIPT_DOCX_OUT = OUT_DIR / "IOS-KOL影片完整逐字稿_含時間碼_20260615.docx"
EXPANDED_DOCX_OUT = OUT_DIR / "AI伺服器供應鏈六大主題_完整版研究稿_20260615.docx"


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")


def configure_doc(doc: Document, font_size: int = 10) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    normal.font.size = Pt(font_size)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=15 if level == 1 else 12, bold=True)
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_para(doc: Document, text: str, size: int = 10, bold: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=8, bold=bold)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[i], "F3F6F8")
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def parse_srt_time(ts: str) -> int:
    hh, mm, rest = ts.split(":")
    ss, ms = rest.split(",")
    return ((int(hh) * 60 + int(mm)) * 60 + int(ss)) * 1000 + int(ms)


def format_hhmmss(ms: int) -> str:
    total = ms // 1000
    hh = total // 3600
    mm = (total % 3600) // 60
    ss = total % 60
    return f"{hh:02d}:{mm:02d}:{ss:02d}"


def parse_srt(path: Path) -> list[dict[str, str | int]]:
    raw = path.read_text(encoding="utf-8", errors="replace").strip()
    blocks = re.split(r"\n\s*\n", raw)
    segments: list[dict[str, str | int]] = []
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) < 3:
            continue
        time_line = next((line for line in lines if "-->" in line), "")
        if not time_line:
            continue
        start_raw, end_raw = [part.strip() for part in time_line.split("-->", 1)]
        text_lines = lines[lines.index(time_line) + 1 :]
        text = " ".join(text_lines).strip()
        if not text:
            continue
        segments.append(
            {
                "start_ms": parse_srt_time(start_raw),
                "end_ms": parse_srt_time(end_raw),
                "start": start_raw.replace(",", "."),
                "end": end_raw.replace(",", "."),
                "text": normalize_asr_text(text),
            }
        )
    return segments


def normalize_asr_text(text: str) -> str:
    replacements = {
        "Cobos": "CoWoS",
        "COPOS": "CoWoS",
        "Core先進封裝": "CoWoS 先進封裝",
        "SOSC": "SoIC",
        "SYC": "SoIC",
        "HBDC": "HVDC",
        "HVDC": "HVDC",
        "BBU": "BBU",
        "TOMAHOG": "Tomahawk",
        "MLCC": "MLCC",
        "FONC": "FOMC",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def write_transcript_txt(segments: list[dict[str, str | int]]) -> None:
    lines = [
        "IOS-KOL 影片完整逐字稿（含時間碼）",
        "來源：YouTube 本地 ASR / whisper-cpp",
        "注意：此為機器轉錄，專有名詞已做部分語意校正，但仍需人工校稿。",
        "",
    ]
    last_bucket = -1
    for seg in segments:
        bucket = int(seg["start_ms"]) // (10 * 60 * 1000)
        if bucket != last_bucket:
            lines.append("")
            lines.append(f"===== {format_hhmmss(bucket * 10 * 60 * 1000)} =====")
            last_bucket = bucket
        lines.append(f"[{seg['start']} - {seg['end']}] {seg['text']}")
    TRANSCRIPT_TXT_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_transcript_docx(segments: list[dict[str, str | int]]) -> None:
    doc = Document()
    configure_doc(doc, font_size=8)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IOS-KOL 影片完整逐字稿（含時間碼）")
    set_run_font(run, size=18, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)
    add_para(doc, "來源：YouTube 本地 ASR / whisper-cpp｜整理日期：2026-06-15", size=9)
    add_para(doc, "注意：此為機器轉錄，專有名詞已做部分語意校正，但仍需人工校稿；投資與公司判讀請以財報、法說與公開文件二次驗證。", size=9, bold=True)
    add_para(doc, f"原始 SRT：{SRT_PATH}", size=8)
    add_para(doc, f"原始 TXT：{TXT_PATH}", size=8)

    last_bucket = -1
    for seg in segments:
        bucket = int(seg["start_ms"]) // (10 * 60 * 1000)
        if bucket != last_bucket:
            add_heading(doc, format_hhmmss(bucket * 10 * 60 * 1000), 1)
            last_bucket = bucket
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        time_run = p.add_run(f"[{seg['start']} - {seg['end']}] ")
        set_run_font(time_run, size=7, bold=True)
        time_run.font.color.rgb = RGBColor(90, 90, 90)
        text_run = p.add_run(str(seg["text"]))
        set_run_font(text_run, size=8)

    doc.core_properties.title = "IOS-KOL 影片完整逐字稿（含時間碼）"
    doc.core_properties.author = "MAPLAB IOS-KOL Influencer Radar Manager"
    doc.save(TRANSCRIPT_DOCX_OUT)


def build_expanded_report() -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 伺服器供應鏈六大主題：完整版研究稿")
    set_run_font(run, size=18, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("截圖主題 + 影片逐字稿 + 產業第二層思考")
    set_run_font(run2, size=12, bold=True)
    add_para(doc, "整理日期：2026-06-15｜用途：產業追蹤、KOL 雷達、下半年題材拆解", size=9)

    add_heading(doc, "1. 這場分享會真正講的是什麼", 1)
    for item in [
        "表面上六個題目分散在半導體廠務、被動元件、先進封裝、光通訊、功率元件與電子零組件，但底層其實同一條主線：AI data center 從 GPU 供給瓶頸，進入 rack/system infrastructure 瓶頸。",
        "第一階段市場只問 GPU 夠不夠；第二階段要問晶片能不能封裝、rack 能不能供電、資料能不能互聯、熱能不能帶走、電力能不能備援、工廠能不能準時蓋出來。",
        "所以本題不是單一概念股整理，而是『瓶頸遷移地圖』：誰掌握瓶頸，誰才有定價與毛利的機會。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 截圖六大主題的加長版判讀", 1)
    add_table(
        doc,
        ["主題", "第一層解讀", "第二層解讀", "第三層查核"],
        [
            ["半導體廠務", "AI/先進封裝擴產帶動無塵室與機電需求。", "真正瓶頸是可準時交付的 cleanroom、MEP、氣體化學、超純水與電力系統；它們決定新產能能否變成可用產能。", "看工程公司在手訂單、毛利率、案場進度、客戶 capex、缺工與材料成本。"],
            ["被動元件", "MLCC、鋁電、電感、電阻被 AI server 拉動。", "重點不是所有被動都漲，而是 PSU/BBU/rack power 對高壓、高溫、低 ESR、高可靠品項形成資格認證門檻。", "看交期、ASP、庫存天數、AI/工控/車用占比、客戶 design-in。"],
            ["先進封裝", "CoWoS/SoIC/3D 封裝支撐 AI chip。", "先進封裝已經是運算架構的一部分，會影響 memory bandwidth、die size、yield 與 power delivery。", "看 CoWoS/SoIC 月產能、設備精度、Hybrid bonding、載板/PCB、測試良率。"],
            ["光通訊", "OFC 觀察指向 800G/1.6T、NPO/CPO、矽光子。", "當 cluster scale-out 放大，光通訊會從連接配角變成效率瓶頸；CPO/NPO 是架構轉移，不只是模組升級。", "看 switch roadmap、CPO/NPO timing、可插拔模組生命周期、光引擎良率。"],
            ["功率元件", "800V HVDC 與電源架構升級。", "高功率 rack 下，電力路徑本身變成 AI infrastructure 主戰場，SST、PowerRack、BBU、PSU、SiC/GaN 都會被重新定義。", "看正負 400V 過渡、800VDC 安規、design win、量產平台、效率/可靠性。"],
            ["電子零組件", "BBU、連接器、散熱、電源、機構件受惠。", "AI server 已從 board-level 競爭變成 rack-level 系統工程，價值會往高功率、高可靠、易維修零組件集中。", "看 25kW BBU、液冷 CDU/cold plate、connector、rack 出貨與客戶驗證。"],
        ],
    )

    add_heading(doc, "3. 逐字稿對應到的主線", 1)
    for item in [
        "開場已明確說明 AI 影響不是只局限 GPU，而是一路延伸到先進製程、先進封裝、電源架構、被動元件、功率元件與光通訊。這是全文判讀的核心句。",
        "先進封裝段落反覆提到 CoWoS、SoIC、載板、PCB 設備、Hybrid bonding 與 CPO 相關材料，表示節目把 AI 供應鏈瓶頸放在封裝/設備/基板三者的交會處。",
        "800V HVDC 段落不是只談電源廠，而是從傳統供電鏈的轉換損耗切入，再延伸到 PowerRack、SST、BBU 與 GPU 供電，這是 rack-level redesign。",
        "BBU 段落提到從選配走向標配、5.5kW 到 25kW、SMT 板與被動元件用量提升，這把被動元件、電源、電子零組件三個章節串起來。",
        "CPO 段落提到 Tomahawk CPO 與玻璃核心載板，代表光通訊不只光模組，而會連到 switch ASIC、封裝、載板與散熱。",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "4. 第二層思考：六條供應鏈怎麼互相牽動", 1)
    for heading, bullets in [
        (
            "4.1 先進製程與廠務不是背景音，是供給開關",
            [
                "如果 2nm/A16/HPC 產能要開出來，廠房、無塵室、電力、冷卻、氣體、化學品與安全系統都要先到位。",
                "因此廠務供應鏈的投資邏輯不是『AI 很熱所以蓋廠』，而是『高階產能的可用性取決於廠務瓶頸』。",
                "驗證時要避免只看營收年增，還要看案場毛利、工期、在手訂單品質與客戶集中度。",
            ],
        ),
        (
            "4.2 先進封裝把設備、材料、載板、測試串成同一個瓶頸",
            [
                "CoWoS 的重點是 HBM 與 accelerator 的高頻寬連接；SoIC/Hybrid bonding 的重點是更高密度與更短 interconnect。",
                "當封裝尺寸變大、精度提高、基板變大，設備資本密集度、良率與週期時間都會改變。",
                "第二層受惠名單應拆成：前段設備、封裝設備、載板/PCB 設備、測試/檢測、材料與工程服務。",
            ],
        ),
        (
            "4.3 800VDC 不是一個零件題材，而是一整條電力架構改造",
            [
                "rack power 越高，低壓大電流造成的銅用量、線損、熱與空間問題越嚴重。",
                "800VDC/正負 400V 的意義是降低電流與損耗，但導入會受安規、標準、客戶平台與過渡方案影響。",
                "因此不能把所有 800V 關鍵字都當成同一個受惠，必須分辨誰有量產認證、誰只是展品或概念。",
            ],
        ),
        (
            "4.4 BBU 是 AI rack 的保險絲，也是被動元件增量入口",
            [
                "當 BBU 從選配走向標配，價值鏈會從電池芯延伸到 BMS、SMT 板、功率元件、連接器、被動元件與整機架認證。",
                "25kW BBU 的價值不只在單價提高，也在於安規、熱設計、可靠性與客戶導入週期。",
                "被動元件的增量若來自 PSU/BBU，比只來自主板更有結構性。",
            ],
        ),
        (
            "4.5 CPO/NPO 會改變光通訊價值分配",
            [
                "可插拔光模組仍會存在，但當頻寬往 1.6T/3.2T 推進，光電轉換距離 switch ASIC 太遠會增加功耗與訊號問題。",
                "CPO/NPO 的價值在於把光引擎拉近運算/交換晶片，但也帶來維修性、熱管理、標準化與良率問題。",
                "所以短期看 800G/1.6T 出貨，中期看平台導入，長期看矽光子與封裝整合能力。",
            ],
        ),
        (
            "4.6 電子零組件要從 server BOM 轉成 rack BOM 來看",
            [
                "傳統看 server 零組件容易分散在電源、散熱、線材、機構、連接器；AI rack 會把這些零組件重新整合成系統能力。",
                "未來高價值會集中在高功率、高可靠、可維修、已通過客戶平台驗證的供應商。",
                "判斷重點是 rack 出貨、platform share、液冷方案、connector 規格與 BBU 滲透率。",
            ],
        ),
    ]:
        add_heading(doc, heading, 2)
        for bullet in bullets:
            add_bullet(doc, bullet)

    add_heading(doc, "5. 觀察名單應該怎麼分層", 1)
    add_table(
        doc,
        ["層級", "類型", "要問的問題", "常見誤判"],
        [
            ["第一層", "直接產品供應商", "是否已經進入 AI server/rack 客戶供應鏈？量產還是送樣？", "看到展品就當量產。"],
            ["第二層", "瓶頸型供應商", "它是否掌握交期、良率、安規或資格認證？", "把低門檻代工也當成高定價瓶頸。"],
            ["第三層", "設備/材料/工程", "它是否受益於客戶 capex 與製程複雜度提高？", "只看 capex 總量，不看產品曝險。"],
            ["第四層", "替代/轉單", "供需失衡是否真的造成新供應商切入？", "把一次性轉單當長期 share gain。"],
        ],
    )

    add_heading(doc, "6. 接下來要追的訊號", 1)
    for item in [
        "TSMC / OSAT 是否持續上修 CoWoS、SoIC、先進封裝設備與封裝廠區投資。",
        "NVIDIA / hyperscaler 是否明確公布 800VDC、正負 400V、PowerRack、SST、BBU 的平台導入節奏。",
        "被動元件是否出現交期拉長、價格回升、AI server 認證進展、庫存天數下降。",
        "CPO/NPO 是否從展會展示進入平台規格，尤其是 switch ASIC、光引擎、雷射、矽光子封裝的量產節點。",
        "BBU 是否從 5.5kW/12kW 往 25kW 滲透，並形成電池芯、BMS、SMT、被動元件、散熱與安規的整體受惠。",
        "廠務工程是否出現案量增加但毛利未被材料與工期侵蝕，否則營收成長未必代表獲利上修。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 風險清單", 1)
    for item in [
        "AI capex 修正：hyperscaler 若放緩建置，所有第二層供應鏈會同步受壓。",
        "導入遞延：800VDC、CPO/NPO、SoIC/Hybrid bonding 都可能因標準、良率、安規或客戶平台延後。",
        "供應商過度擴產：題材初期供不應求，後期若產能一次開出，ASP 與毛利可能回落。",
        "公司名單錯配：同一產業內不是每家公司都受惠，必須確認產品線、客戶、認證與出貨階段。",
        "ASR 風險：逐字稿是機器轉錄，專有名詞與數字需要人工二校；本文把它當方向材料，不當最終精確引用。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. 附錄：原始材料", 1)
    add_bullet(doc, f"完整逐字稿 Word：{TRANSCRIPT_DOCX_OUT}")
    add_bullet(doc, f"完整逐字稿 TXT：{TRANSCRIPT_TXT_OUT}")
    add_bullet(doc, f"原始 ASR TXT：{TXT_PATH}")
    add_bullet(doc, f"原始 ASR SRT：{SRT_PATH}")
    if SCREENSHOT_PATH.exists():
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_heading(doc, "附錄：使用者截圖", 1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(SCREENSHOT_PATH), width=Inches(3.0))

    doc.core_properties.title = "AI 伺服器供應鏈六大主題：完整版研究稿"
    doc.core_properties.author = "MAPLAB IOS-KOL Influencer Radar Manager"
    doc.save(EXPANDED_DOCX_OUT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    segments = parse_srt(SRT_PATH)
    write_transcript_txt(segments)
    build_transcript_docx(segments)
    build_expanded_report()
    print(TRANSCRIPT_TXT_OUT)
    print(TRANSCRIPT_DOCX_OUT)
    print(EXPANDED_DOCX_OUT)
    print(f"segments={len(segments)}")


if __name__ == "__main__":
    main()
